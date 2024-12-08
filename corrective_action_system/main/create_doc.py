# create_doc.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add header information
header_text = [
    "Republic of the Philippines",
    "SORSOGON STATE UNIVERSITY",
    "Magsaysay Street, Salog (Pob.), Sorsogon City, Sorsogon",
    "Tel No.: 056 211-0103; Email Add.: qa.iso@sorsu.edu.ph"
]
for line in header_text:
    header_para = doc.add_paragraph(line)
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a title for the form
title = doc.add_heading('REQUEST FOR ACTION', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Create the first part of the table
doc.add_heading('PART 1: What is wrong?', level=2)

# Prompt the user for inputs
originator_name = input("Enter the Originator’s Name/ID No.: ")
unit_department = input("Enter the Unit/Department: ")
phone_number = input("Enter the Phone number: ")
rfa_ref_no = input("Enter the RFA Ref. No. (Auditor): ")
date_issued = input("Enter the Date Issued: ")
rfa_intent = input("Enter the RFA intent (e.g., correct a NC / prevent a potential NC): ")
nc_department = input("Enter the Department where NC exists: ")
nc_related_options = input("Enter the NC-related options (e.g., IQA-Related, Supplier-Related): ")

# Create the table for the non-conformity (NC) data
table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'

# Set the widths and content for each cell
table.cell(0, 0).merge(table.cell(0, 4)).text = 'NON-CONFORMITY (NC) DATA'
table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table.cell(0, 0).paragraphs[0].runs[0].bold = True

# Set row 1 content
table.cell(1, 0).text = "Originator’s Name/ID No."
table.cell(1, 1).text = "Unit/Department"
table.cell(1, 2).text = "Phone"
table.cell(1, 3).text = "RFA Ref. No. (Auditor)"
table.cell(1, 4).text = "Date Issued"

# Fill in the user inputs in row 1
table.cell(1, 0).text = originator_name
table.cell(1, 1).text = unit_department
table.cell(1, 2).text = phone_number
table.cell(1, 3).text = rfa_ref_no
table.cell(1, 4).text = date_issued

# Set row 2 content
table.cell(2, 0).merge(table.cell(2, 2)).text = "This RFA is intended to:"
table.cell(2, 3).merge(table.cell(2, 4)).text = "Department (where NC exists)"

# Fill in the user inputs in row 2
table.cell(2, 0).merge(table.cell(2, 2)).text += f" {rfa_intent}"
table.cell(2, 3).merge(table.cell(2, 4)).text = nc_department

# Add checkboxes for options in row 3
row3 = table.rows[3]
row3.cells[0].text = (
    "□ correct a NC / eliminate source of non-conformance\n"
    "□ prevent a potential NC / mitigate risk\n"
    "□ For Improvement"
)
row3.cells[3].merge(row3.cells[4]).text = (
    f"□ {nc_related_options}\n"
    "□ Process/Procedural-related\n"
    "□ Customer Satisfaction Related\n"
    "□ HRD-Related\n"
    "□ Relates to KPI/Quality Objective Review"
)

# Save the document
file_name = input("Enter the file name to save as (with .docx extension): ")
doc.save(file_name)

print(f"Document saved as {file_name}")
